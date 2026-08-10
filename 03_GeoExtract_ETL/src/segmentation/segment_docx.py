from docx import Document
from src.segmentation.utils import find_raw_file

def segment_docx(doc_id: str):
    """
    DOCX сегментатор + пост-обработка в одном блоке.
    Полностью совместим с Golden Set.
    """
    path = find_raw_file(doc_id)
    doc = Document(path)

    blocks = []

    # --- Параграфы и списки ---
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        style = p.style.name

        if style.startswith("Heading"):
            blocks.append({"type": "header", "text": text})
            continue

        if style in ("List Paragraph", "List Bullet", "List Number"):
            blocks.append({"type": "list", "items": [text]})
            continue

        blocks.append({"type": "paragraph", "text": text})

    # --- Таблицы ---
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(cells)

        blocks.append({
            "type": "table",
            "rows": rows,
            "columns": rows[0] if rows else []
        })

    # --- Пост-обработка ---
    out = []
    for b in blocks:
        t = b["type"]
        txt = b.get("text", "").strip()

        # удаляем слишком короткие параграфы
        if t == "paragraph" and len(txt) < 3:
            continue

        # объединяем короткие параграфы (<40 символов)
        if (
            out
            and t == "paragraph"
            and out[-1]["type"] == "paragraph"
            and len(txt) < 40
        ):
            out[-1]["text"] += " " + txt
            continue

        out.append(b)

    return out